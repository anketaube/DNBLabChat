import streamlit as st
import pandas as pd
import httpx
from parsel import Selector
import re
from urllib.parse import urljoin

# NEU: Für PDF/Word
import fitz  # PyMuPDF
from docx import Document

# NEU: LlamaIndex
from llama_index.core import Document as LlamaDocument, VectorStoreIndex

# Basis-URL und Wunsch-URLs
BASE_URL = "https://www.dnb.de/DE/Professionell/Services/WissenschaftundForschung/DNBLab"
START_URL = BASE_URL + "/dnblab_node.html"
EXTRA_URLS = [
    "https://www.dnb.de/DE/Professionell/Services/WissenschaftundForschung/DNBLab/dnblabTutorials.html?nn=849628",
    "https://www.dnb.de/DE/Professionell/Services/WissenschaftundForschung/DNBLabPraxis/dnblabPraxis_node.html",
    "https://www.dnb.de/DE/Professionell/Services/WissenschaftundForschung/DNBLab/dnblabSchnittstellen.html?nn=849628",
    "https://www.dnb.de/DE/Professionell/Services/WissenschaftundForschung/DNBLab/dnblabFreieDigitaleObjektsammlung.html?nn=849628"
]

st.sidebar.title("Konfiguration")
data_source = st.sidebar.radio("Datenquelle wählen:", ["Excel-Datei", "DNBLab-Webseite", "PDF/Word-Dateien"])
chatgpt_model = st.sidebar.selectbox(
    "ChatGPT Modell wählen",
    options=["gpt-3.5-turbo", "gpt-4-turbo"],
    index=1
)
st.sidebar.markdown(f"Verwendetes Modell: **{chatgpt_model}**")

if "OPENAI_API_KEY" not in st.secrets:
    st.error("API-Key fehlt. Bitte in den Streamlit-Secrets hinterlegen.")
    st.stop()
api_key = st.secrets["OPENAI_API_KEY"]

@st.cache_data(show_spinner=True)
def crawl_dnblab():
    client = httpx.Client(timeout=10, follow_redirects=True)
    visited = set()
    to_visit = set([START_URL] + EXTRA_URLS)
    data = []
    while to_visit:
        url = to_visit.pop()
        if url in visited:
            continue
        visited.add(url)
        try:
            resp = client.get(url)
            if resp.status_code != 200:
                continue
            selector = Selector(resp.text)
            content = ' '.join(selector.xpath(
                '//main//text() | //div[@role="main"]//text() | //body//text() | //article//text() | //section//text() | //p//text() | //li//text()'
            ).getall())
            content = re.sub(r'\s+', ' ', content).strip()
            if content and len(content) > 50:
                data.append({
                    'datensetname': f"Web-Inhalt: {url}",
                    'volltextindex': content,
                    'quelle': url
                })
            # Interne Links sammeln
            for link in selector.xpath('//a/@href').getall():
                full_url = urljoin(url, link).split('#')[0]
                if (
                    full_url.startswith(BASE_URL)
                    and full_url not in visited
                    and full_url not in to_visit
                ):
                    to_visit.add(full_url)
        except Exception as e:
            st.warning(f"Fehler beim Crawlen von {url}: {e}")
    if data:
        df = pd.DataFrame(data)
        df.columns = df.columns.str.strip().str.lower()
        return df
    else:
        return None

@st.cache_data
def load_excel(file):
    try:
        xls = pd.ExcelFile(file)
        df = pd.read_excel(xls, sheet_name=xls.sheet_names[0], header=0, na_filter=False)
        df['volltextindex'] = df.apply(lambda row: ' | '.join(str(cell) for cell in row if pd.notnull(cell)), axis=1)
        df['quelle'] = "Excel-Datei"
        df.columns = df.columns.str.strip().str.lower()
        return df
    except Exception as e:
        st.error(f"Fehler beim Laden der Excel-Datei: {e}")
        return None

def load_pdf(file):
    try:
        doc = fitz.open(stream=file.read(), filetype="pdf")
        text = " ".join(page.get_text() for page in doc)
        return pd.DataFrame({'volltextindex': [text], 'quelle': [file.name]})
    except Exception as e:
        st.error(f"Fehler beim Laden der PDF-Datei: {e}")
        return None

def load_word(file):
    try:
        doc = Document(file)
        text = " ".join(para.text for para in doc.paragraphs)
        return pd.DataFrame({'volltextindex': [text], 'quelle': [file.name]})
    except Exception as e:
        st.error(f"Fehler beim Laden der Word-Datei: {e}")
        return None

def df_to_documents(df):
    # Wandelt DataFrame in LlamaIndex-Dokumente mit Metadaten um
    return [
        LlamaDocument(text=row["volltextindex"], metadata={"quelle": row["quelle"]})
        for _, row in df.iterrows()
    ]

st.title("DNBLab-Chatbot")

df = None

if data_source == "Excel-Datei":
    uploaded_file = st.sidebar.file_uploader("Excel-Datei hochladen", type=["xlsx"])
    if uploaded_file:
        with st.spinner("Excel-Datei wird geladen..."):
            df = load_excel(uploaded_file)
elif data_source == "DNBLab-Webseite":
    st.sidebar.info("Es wird die DNBLab-Webseite inkl. aller Unterseiten indexiert. Das kann einige Sekunden dauern.")
    with st.spinner("DNBLab-Webseite wird indexiert..."):
        df = crawl_dnblab()
elif data_source == "PDF/Word-Dateien":
    uploaded_files = st.sidebar.file_uploader("PDF/Word-Dateien hochladen", type=["pdf", "docx"], accept_multiple_files=True)
    dfs = []
    if uploaded_files:
        with st.spinner("Dateien werden geladen..."):
            for file in uploaded_files:
                if file.name.endswith(".pdf"):
                    tdf = load_pdf(file)
                elif file.name.endswith(".docx"):
                    tdf = load_word(file)
                else:
                    tdf = None
                if tdf is not None:
                    dfs.append(tdf)
            if dfs:
                df = pd.concat(dfs, ignore_index=True)

if df is None or df.empty:
    st.info("Bitte laden Sie eine Datei hoch oder wählen Sie die DNBLab-Webseite aus der Sidebar.")
    st.stop()

st.write(f"Geladene Datensätze: {len(df)}")
st.markdown("**Folgende Quellen wurden indexiert:**")
for url in sorted(df['quelle'].unique()):
    st.markdown(f"- {url}")

# LlamaIndex: Index aufbauen
with st.spinner("Index wird erstellt..."):
    documents = df_to_documents(df)
    index = VectorStoreIndex.from_documents(documents)
    query_engine = index.as_query_engine(similarity_top_k=5)

query = st.text_input("Suchbegriff oder Frage eingeben:")

if query:
    with st.spinner("Frage wird analysiert..."):
        response = query_engine.query(query)
        antwort_text = response.response
        quellen = [n.metadata["quelle"] for n in response.source_nodes]

    st.subheader("Antwort des Sprachmodells:")
    st.write(antwort_text)
    st.markdown("**Quellen:**")
    for quelle in set(quellen):
        st.markdown(f"- {quelle}")

    # Optional: Treffer-DataFrame anzeigen
    if len(quellen) > 0:
        treffer = df[df['quelle'].isin(quellen)]
        if not treffer.empty:
            st.subheader("Relevante Treffer aus den Daten:")
            st.dataframe(treffer)
else:
    st.info("Bitte geben Sie einen Suchbegriff oder eine Frage ein.")

